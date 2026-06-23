"""Generate the MaxCode Pipeline Technical Reference as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ── Title ──
title = doc.add_heading("MaxCode Migration Pipeline", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Technical Reference — PyTorch to JAX/Flax Conversion")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# 1. Overview
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1. Pipeline Overview", level=1)
doc.add_paragraph(
    "MaxCode converts PyTorch repositories to JAX/Flax through a five-step "
    "pipeline. Each step is an independent script that reads the output of "
    "the previous step, allowing re-runs without restarting from scratch."
)

# Steps table
table = doc.add_table(rows=6, cols=3, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Step", "Script", "Purpose"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

steps = [
    ("1 — Clone", "step1_clone_repo.py",
     "Fetch the PyTorch repository from GitHub"),
    ("2 — Index", "step2_populate_rag.py",
     "Build the RAG vector database from reference JAX/Flax sources"),
    ("3 — Merge", "step3_merge.py",
     "Auto-detect model files AND utility files, resolve dependencies, "
     "merge into two files (model + utilities)"),
    ("4 — Convert", "step4_convert.py",
     "Convert both model and utility files with RAG context, fill gaps, "
     "validate, and repair"),
    ("5 — Verify", "step5_verify.py",
     "Score completeness (AST) and correctness (LLM) of model and utility output"),
]
for row_idx, (step, script, purpose) in enumerate(steps, 1):
    table.rows[row_idx].cells[0].text = step
    table.rows[row_idx].cells[1].text = script
    table.rows[row_idx].cells[2].text = purpose

doc.add_paragraph()
doc.add_paragraph(
    "The pipeline produces two JAX/Flax output files: one for model "
    "definitions (nn.Module subclasses) and one for utility/helper code "
    "(custom ops, persistence, misc functions). This two-file approach "
    "gives the LLM full context within each domain while ensuring the "
    "output is self-contained with no broken imports."
)

# ── Key output files ──
doc.add_heading("1.1 Key Artefacts", level=2)
t_artefacts = doc.add_table(rows=7, cols=2, style="Light Shading Accent 1")
t_artefacts.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["File", "Description"]):
    t_artefacts.rows[0].cells[i].text = h
    for r in t_artefacts.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
artefacts = [
    ("merged_model.py", "All nn.Module files merged in dependency order (Step 3)"),
    ("merged_utils.py", "All transitively-imported utility files merged in "
     "dependency order (Step 3b)"),
    ("output/<repo>_jax.py", "Converted JAX/Flax model code (Step 4)"),
    ("output/<repo>_utils_jax.py", "Converted JAX utility code (Step 4)"),
    ("output/verification_scorecard.json", "Completeness and correctness "
     "scores for both model and utility output (Step 5)"),
    ("~/rag_store.db", "SQLite vector database with embedded reference "
     "documents (Step 2)"),
]
for row_idx, (f, d) in enumerate(artefacts, 1):
    t_artefacts.rows[row_idx].cells[0].text = f
    t_artefacts.rows[row_idx].cells[1].text = d


# ══════════════════════════════════════════════════════════════════════
# 2. Configuration
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("2. Configuration (config.py)", level=1)
doc.add_paragraph(
    "All paths, filtering rules, and helper functions live in config.py. "
    "Scripts import what they need so every setting has a single source of truth."
)

t_cfg = doc.add_table(rows=9, cols=2, style="Light Shading Accent 1")
t_cfg.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Constant", "Purpose"]):
    t_cfg.rows[0].cells[i].text = h
    for r in t_cfg.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
cfg_rows = [
    ("REPO_URL / REPO_DIR", "Target repository URL and local clone path"),
    ("MERGED_FILE", "Path to merged_model.py (model merge output)"),
    ("MERGED_UTILS_FILE", "Path to merged_utils.py (utility merge output)"),
    ("OUTPUT_DIR", "Directory for converted JAX files and scorecard"),
    ("RAG_SOURCE_DIR", "Directory of reference .py files for the RAG database"),
    ("MERGE_EXCLUDE_PATHS", "Glob patterns to exclude from model merge "
     "(e.g. megatron/model/fused_*.py)"),
    ("MERGE_EXCLUDE_CLASSES", "Class name patterns to exclude from model merge "
     "(e.g. *Pipe, ColumnParallelLinear)"),
    ("MERGE_EXCLUDE_UTILS", "Glob patterns to exclude from utility merge "
     "(setup.py, test files, etc.)"),
]
for row_idx, (c, p) in enumerate(cfg_rows, 1):
    t_cfg.rows[row_idx].cells[0].text = c
    t_cfg.rows[row_idx].cells[1].text = p


# ══════════════════════════════════════════════════════════════════════
# 3. Step 1 — Clone
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("3. Repository Cloning (Step 1)", level=1)
doc.add_paragraph(
    "step1_clone_repo.py accepts an optional repository URL on the command "
    "line, persists it to .repo_url for subsequent steps, and runs git clone. "
    "If the directory already exists it skips cloning. After cloning it walks "
    "the directory tree and prints a summary of Python file and line counts."
)


# ══════════════════════════════════════════════════════════════════════
# 4. Step 2 — RAG Indexing
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("4. RAG Indexing Strategy (Step 2)", level=1)

doc.add_heading("4.1 Document Corpus", level=2)
doc.add_paragraph(
    "The RAG database contains 48 reference documents stored under "
    "MaxCode/rag/sources/, split into two categories:"
)
doc.add_paragraph(
    "Generic references (24 files) — JAX/Flax API documentation, MaxText "
    "model implementations, Flash-linear-attention examples, Flax attention "
    "patterns.",
    style="List Bullet",
)
doc.add_paragraph(
    "Targeted patterns (24 files) — WRONG/CORRECT/WHY triplets covering "
    "common conversion mistakes: incorrect cosine similarity, wrong einsum "
    "dimensions, missing weight initialisation, broken MoE routing, etc.",
    style="List Bullet",
)

doc.add_heading("4.2 Embedding Flow", level=2)
doc.add_paragraph(
    "Each .py file in the source directory goes through the following pipeline:"
)
for item in [
    "Read the file content.",
    "Generate a structured description using Gemini (CODE_DESCRIPTION prompt) "
    "that captures the file's functionality and usage in JSON format.",
    "Embed the description (not the raw code) using Google's embedding-001 "
    "model. This produces a dense vector in float32.",
    "Store the document in a SQLite database (rag_store.db) with columns: "
    "id, name, text (full source), desc (generated description), file (path), "
    "embedding (pickled numpy array).",
]:
    doc.add_paragraph(item, style="List Number")

doc.add_paragraph(
    "A 2-second sleep is enforced between embedding API calls to respect "
    "rate limits. Results are cached in-memory to avoid redundant calls "
    "within the same session."
)

doc.add_heading("4.3 Vector Index", level=2)
doc.add_paragraph(
    "At query time, all stored embeddings are loaded into a NumPy array "
    "(shape: num_docs x embedding_dim). Search uses squared L2 (Euclidean) "
    "distance with np.argsort to find the top-k nearest neighbours. There "
    "is no approximate nearest-neighbour index (FAISS, Annoy, etc.) — the "
    "corpus is small enough (~48 docs) for exact brute-force search."
)

doc.add_heading("4.4 Key Parameters", level=2)
t2 = doc.add_table(rows=7, cols=3, style="Light Shading Accent 1")
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Parameter", "Value", "Location"]):
    t2.rows[0].cells[i].text = h
    for r in t2.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for row_idx, (p, v, loc) in enumerate([
    ("Embedding model", "models/embedding-001 (Google)", "embedding.py"),
    ("Description model", "Gemini 2.5 Flash", "step2_populate_rag.py"),
    ("Distance metric", "Squared L2 (Euclidean)", "vector_db.py"),
    ("Storage format", "SQLite + pickled float32 arrays", "vector_db.py"),
    ("API sleep", "2 seconds between calls", "embedding.py"),
    ("Max context length", "100,000 characters", "rag_agent.py"),
], 1):
    t2.rows[row_idx].cells[0].text = p
    t2.rows[row_idx].cells[1].text = v
    t2.rows[row_idx].cells[2].text = loc


# ══════════════════════════════════════════════════════════════════════
# 5. Step 3 — Merge (Model + Utilities)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("5. Merge Strategy (Step 3)", level=1)
doc.add_paragraph(
    "Step 3 has two phases: Step 3a merges model files (nn.Module "
    "subclasses) into merged_model.py, and Step 3b discovers and merges "
    "transitively-imported utility files into merged_utils.py."
)

# -- 5.1 Model File Detection --
doc.add_heading("5.1 Model File Detection (Step 3a)", level=2)
doc.add_paragraph(
    "The merge script scans every .py file in the repository and identifies "
    "model files by parsing the AST looking for class definitions that "
    "subclass nn.Module (matching torch.nn.Module, nn.Module, or bare Module). "
    "Files are opened with utf-8-sig encoding to handle BOM characters."
)

# -- 5.2 File-Level Filtering --
doc.add_heading("5.2 File-Level Filtering", level=2)
doc.add_paragraph("Before merging, several file-level filters are applied:")
for f in [
    "Config exclude patterns — path globs defined in config.py "
    "(MERGE_EXCLUDE_PATHS).",
    "Fused kernel heuristic — files matching fused_*.py are skipped.",
    "Infrastructure files — files where every class subclasses an infrastructure "
    "base (autograd.Function, PipelineModule, TransformerEngine wrappers, Enum) "
    "AND the file imports infrastructure packages (apex, deepspeed, "
    "transformer_engine).",
]:
    doc.add_paragraph(f, style="List Bullet")

# -- 5.3 Dependency Resolution --
doc.add_heading("5.3 Dependency Resolution", level=2)
doc.add_paragraph(
    "An import graph is built between the remaining model files by parsing "
    "ImportFrom AST nodes and resolving them to file paths (both relative "
    "and absolute-style imports). Entry points are identified as files that "
    "are not imported by any other model file but do import at least one. "
    "A BFS + DFS post-order traversal produces a topological ordering: "
    "dependencies first, entry points last."
)

# -- 5.4 Model Merge Process --
doc.add_heading("5.4 Model Merge Process", level=2)
for item in [
    "Standard-library imports are de-duplicated and collected at the top.",
    "Local cross-file imports are removed (no longer needed in a single file).",
    "Empty blocks left behind by import removal get a 'pass' statement inserted.",
    "Code sections are concatenated with file-boundary comments.",
    "A second pass removes infrastructure classes from the merged output "
    "(autograd.Function subclasses, PipelineModule, TransformerEngine wrappers, "
    "Enum subclasses, *Pipe-suffixed classes).",
]:
    doc.add_paragraph(item, style="List Number")

doc.add_paragraph(
    "The result is merged_model.py with all model definitions in dependency "
    "order, ready for conversion."
)

# -- 5.5 Utility File Discovery (Step 3b) --
doc.add_heading("5.5 Utility File Discovery (Step 3b)", level=2)
doc.add_paragraph(
    "After the model merge, Step 3b discovers all Python files transitively "
    "imported by model files within the same repository. This ensures the "
    "converted output is self-contained — no broken imports referencing "
    "modules that were never converted."
)

doc.add_heading("Discovery: BFS from Model Files", level=3)
doc.add_paragraph(
    "Starting from the final set of model files included in the merge, "
    "find_all_local_dependencies() performs a breadth-first search through "
    "all local imports (using the same get_local_imports() parser that "
    "handles the model import graph). Every transitively-reachable .py "
    "file within the repository is collected. Files already in the model "
    "set are excluded — only non-model utility files are returned."
)

doc.add_heading("Classification", level=3)
doc.add_paragraph(
    "Each discovered utility file is classified by classify_utility_file() "
    "into one of five categories:"
)

t_cat = doc.add_table(rows=6, cols=3, style="Light Shading Accent 1")
t_cat.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Category", "Detection", "Action"]):
    t_cat.rows[0].cells[i].text = h
    for r in t_cat.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
cats = [
    ("init_reexport",
     "__init__.py whose body only contains imports, assignments, and "
     "docstrings (re-export files)",
     "Skip — content is inlined by the merge"),
    ("cuda_kernel",
     "Files that call load() or load_inline() AND reference .cu or .cpp "
     "files (CUDA plugin loaders)",
     "Skip — no JAX equivalent for custom CUDA kernels"),
    ("torch_autograd",
     "Files with classes subclassing torch.autograd.Function",
     "Keep — these typically have a Python fallback path worth converting"),
    ("torch_utility",
     "Files that import torch or torch.* modules",
     "Keep — PyTorch-dependent utility code to convert"),
    ("pure_python",
     "Files with no torch dependency",
     "Keep — pure Python helpers, data structures, etc."),
]
for row_idx, (cat, detect, action) in enumerate(cats, 1):
    t_cat.rows[row_idx].cells[0].text = cat
    t_cat.rows[row_idx].cells[1].text = detect
    t_cat.rows[row_idx].cells[2].text = action

doc.add_heading("Filtering", level=3)
doc.add_paragraph(
    "Before classification, utility files are checked against "
    "MERGE_EXCLUDE_UTILS glob patterns (setup.py, test files, etc.). "
    "After classification, init_reexport and cuda_kernel files are removed. "
    "The function returns the kept files, removed files with reasons, and "
    "a category map."
)

doc.add_heading("Ordering and Merging", level=3)
doc.add_paragraph(
    "The kept utility files are topologically sorted by their internal "
    "import graph (same DFS post-order algorithm as the model merge). "
    "They are then merged into merged_utils.py using the same merge_files() "
    "function: imports deduplicated, local imports removed, empty blocks "
    "fixed. The utility merge is kept separate from the model merge to "
    "avoid mixing concerns."
)

# -- 5.6 Example output --
doc.add_heading("5.6 Example: stylegan2-ada-pytorch", level=2)
doc.add_paragraph(
    "For the stylegan2-ada-pytorch repository, Step 3b discovers and "
    "processes the following utility files:"
)
doc.add_paragraph(
    "Discovered and kept: torch_utils/misc.py, torch_utils/persistence.py, "
    "torch_utils/ops/bias_act.py, torch_utils/ops/upfirdn2d.py, "
    "torch_utils/ops/conv2d_resample.py, torch_utils/ops/fma.py, "
    "dnnlib/util.py",
    style="List Bullet",
)
doc.add_paragraph(
    "Filtered out: torch_utils/ops/custom_ops.py (CUDA kernel loader), "
    "various __init__.py files (re-exports)",
    style="List Bullet",
)
doc.add_paragraph(
    "Without Step 3b, the converted model output would have broken imports "
    "referencing misc, bias_act, conv2d_resample, upfirdn2d, fma, and "
    "dnnlib — modules that were never converted.",
    style="List Bullet",
)


# ══════════════════════════════════════════════════════════════════════
# 6. Retrieval Strategy
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("6. Retrieval Strategy", level=1)

doc.add_heading("6.1 Hybrid Per-Component Retrieval", level=2)
doc.add_paragraph(
    "Both conversion agents (SingleFileAgent, ModelConversionAgent) use "
    "retrieve_per_component_context(), which combines two strategies:"
)

doc.add_heading("Full-File Query (Broad Context)", level=3)
doc.add_paragraph(
    "The entire PyTorch source code is embedded as a single query and "
    "the top 15 results are retrieved. This captures the overall domain "
    "(transformer architecture, attention patterns, etc.) and provides "
    "broad reference material."
)

doc.add_heading("Per-Component Queries (Targeted Context)", level=3)
doc.add_paragraph(
    "The source code is parsed with Python's ast module to extract each "
    "top-level class and function. A focused query string is built for each:"
)
doc.add_paragraph(
    'Classes: "JAX Flax {ClassName} {base_classes} {method_names} {init_params}"',
    style="List Bullet",
)
doc.add_paragraph(
    'Functions: "JAX Flax {func_name} {param_names}"',
    style="List Bullet",
)
doc.add_paragraph(
    "If there are more than 12 components, signatures are batched in groups "
    "of 4 to cap the number of embedding API calls at roughly 3-5."
)

doc.add_heading("Deduplication and Ranking", level=3)
doc.add_paragraph(
    "Results from both the full-file query and all per-component queries "
    "are merged into a single set, deduplicated by file path (keeping the "
    "entry with the best distance for each file). The final list is sorted "
    "by distance and truncated to max_total (default 15). If AST parsing "
    "fails, the method falls back to a single full-file query."
)

t3 = doc.add_table(rows=4, cols=2, style="Light Shading Accent 1")
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Parameter", "Default"]):
    t3.rows[0].cells[i].text = h
    for r in t3.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for row_idx, (p, v) in enumerate([
    ("top_k_per_component", "3"),
    ("max_total", "15"),
    ("Batch threshold", ">12 components"),
], 1):
    t3.rows[row_idx].cells[0].text = p
    t3.rows[row_idx].cells[1].text = v


# ══════════════════════════════════════════════════════════════════════
# 7. Conversion Pipeline (Step 4)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("7. Conversion Pipeline (Step 4)", level=1)

doc.add_heading("7.1 Model Selection", level=2)
doc.add_paragraph(
    "Step 4 initialises a PrimaryAgent and probes available Gemini models "
    "in preference order: Gemini 3.1 Pro Preview, Gemini 2.5 Pro, "
    "Gemini 2.5 Flash. The first model that responds successfully is used "
    "for all conversion and gap-filling calls."
)

doc.add_heading("7.2 Agent Routing", level=2)
doc.add_paragraph(
    "The PrimaryAgent receives the merged file path and orchestrates "
    "the conversion. For each file, it decides which specialised agent "
    "to use:"
)
doc.add_paragraph(
    "ModelConversionAgent — for files containing nn.Module subclasses "
    "(detected by is_model_file()). Uses MODEL_CONVERSION_PROMPT with "
    "16 conversion rules covering @nn.compact, KV caches, MoE dispatch, "
    "fused QKV projections, float32 softmax upcast, etc.",
    style="List Bullet",
)
doc.add_paragraph(
    "SingleFileAgent — for utility code, training loops, and data loading. "
    "Uses MIGRATE_MODULE_TO_JAX_PROMPT with general JAX best practices.",
    style="List Bullet",
)
doc.add_paragraph(
    "Both agents inject RAG context (retrieved via the hybrid strategy above) "
    "directly into the prompt alongside the PyTorch source code."
)

doc.add_heading("7.3 Model Conversion", level=2)
doc.add_paragraph(
    "The merged_model.py file is passed to PrimaryAgent.run() which routes "
    "it to the ModelConversionAgent. The agent retrieves per-component RAG "
    "context, builds a prompt with the source and reference patterns, and "
    "calls the Gemini LLM. The response is stripped of markdown formatting."
)

doc.add_heading("7.4 Gap-Filling (Two Phases)", level=2)
doc.add_paragraph(
    "After the initial conversion, _fill_missing_components() runs two "
    "phases to catch what the LLM missed:"
)

doc.add_heading("Phase 1 — Missing Top-Level Components", level=3)
doc.add_paragraph(
    "An AST diff compares class and function names between the PyTorch "
    "source and the JAX output. Any top-level component present in the "
    "source but absent in the output is extracted, sent to the LLM with "
    "RAG context, and the converted result is appended to the JAX file."
)

doc.add_heading("Phase 2 — Stub Detection and Missing Methods", level=3)
doc.add_paragraph("Two checks run on the JAX output:")
doc.add_paragraph(
    "Stub detection — walks the AST looking for functions/methods with "
    "placeholder bodies: pass, return None, ... (Ellipsis), docstring-only, "
    "or raise NotImplementedError.",
    style="List Bullet",
)
doc.add_paragraph(
    "Missing-method detection — for each class that exists in both source "
    "and output, compares method sets and identifies methods present in "
    "the PyTorch class but absent from the JAX class.",
    style="List Bullet",
)
doc.add_paragraph(
    "The PyTorch source for all identified stubs and missing methods is "
    "collected and sent in a single LLM call (FILL_STUBS_PROMPT) that "
    "receives the complete JAX file and returns the complete file with "
    "stubs replaced by real implementations. The result is accepted only "
    "if it passes ast.parse() and is at least 50% the length of the original."
)

doc.add_heading("7.5 Utility Conversion", level=2)
doc.add_paragraph(
    "If merged_utils.py exists (produced by Step 3b), it is converted "
    "separately using the SingleFileAgent — not the ModelConversionAgent, "
    "because utility files contain no nn.Module subclasses. The same "
    "two-phase gap-filling (_fill_missing_components) is applied to the "
    "utility output."
)
doc.add_paragraph(
    "The utility conversion is intentionally separate from the model "
    "conversion for two reasons:",
)
doc.add_paragraph(
    "Different agent: utility code needs general JAX migration rules, "
    "not Flax nn.Module conversion rules.",
    style="List Bullet",
)
doc.add_paragraph(
    "Additive design: the model conversion path is unchanged — utility "
    "handling is a new parallel track that cannot break existing behaviour.",
    style="List Bullet",
)

doc.add_heading("7.6 Markdown Stripping", level=2)
doc.add_paragraph(
    "All LLM responses pass through _strip_markdown_formatting() which "
    "extracts the first Python code block from markdown-formatted output. "
    "It handles three cases: (1) properly fenced ```python...``` blocks, "
    "(2) truncated responses where the opening ``` is present but the "
    "closing ``` is missing (common with long outputs), and "
    "(3) triple-quote wrappers."
)


# ══════════════════════════════════════════════════════════════════════
# 8. Validation and Repair Loop
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("8. Validation and Repair Loop", level=1)

doc.add_heading("8.1 Validation Agent", level=2)
doc.add_paragraph(
    "The ValidationAgent performs an LLM-based comparison between the "
    "original PyTorch source and the converted JAX output. It checks "
    "six categories of deviations:"
)

t4 = doc.add_table(rows=7, cols=3, style="Light Shading Accent 1")
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Category", "What It Catches", "Example"]):
    t4.rows[0].cells[i].text = h
    for r in t4.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for row_idx, (cat, what, ex) in enumerate([
    ("default_value", "Constructor parameter defaults changed",
     "init_method changed from xavier_normal to normal(0.02)"),
    ("initialization", "Weight initialisation added or changed",
     "zeros_init added where PyTorch uses default"),
    ("missing_component", "Classes, functions, methods, constants absent",
     "mup_reinitialize_weights method missing from class"),
    ("reduction_op", ".mean() vs .sum() or axis changes",
     "loss.mean() changed to loss.sum()"),
    ("method_placement", "Methods moved between classes or inlined",
     "helper moved from ClassA to ClassB"),
    ("dropped_feature", "Features removed entirely",
     "Sinkhorn error tracking loop removed"),
], 1):
    t4.rows[row_idx].cells[0].text = cat
    t4.rows[row_idx].cells[1].text = what
    t4.rows[row_idx].cells[2].text = ex

doc.add_paragraph()
doc.add_paragraph(
    "Each deviation is assigned a severity (high, medium, or low) and "
    "includes source_snippet, output_snippet, corrected_snippet, and a "
    "fix instruction. The output is a JSON array."
)

doc.add_heading("8.2 Repair Loop", level=2)
doc.add_paragraph(
    "The PrimaryAgent runs up to 3 iterations of validate-then-repair:"
)
for item in [
    "Validate: run the ValidationAgent to produce a list of deviations.",
    "Exit early if zero deviations remain (clean).",
    "Exit early if deviation count did not decrease from the previous "
    "iteration (no progress — avoid infinite loops).",
    "Filter actionable deviations: skip any whose fix text contains "
    "phrases like 'not recommended', 'desirable deviation', or 'acceptable'.",
    "Build repair prompt: inject the original PyTorch source, current JAX "
    "code, formatted deviation blocks, and RAG context (top 15 results "
    "queried from deviation categories and fix descriptions).",
    "The LLM returns the complete repaired JAX file. Accept only if the "
    "result is at least 50% the length of the input.",
]:
    doc.add_paragraph(item, style="List Number")

doc.add_paragraph(
    "After the loop completes, validation results are stored per file "
    "with full iteration history (deviation counts per iteration, "
    "initial and remaining deviations)."
)


# ══════════════════════════════════════════════════════════════════════
# 9. Verification Scorecard (Step 5)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("9. Verification Scorecard (Step 5)", level=1)

doc.add_heading("9.1 Completeness Score (AST-Based, No LLM)", level=2)
doc.add_paragraph(
    "Both the source and output files are parsed with Python's ast module. "
    "Three component types are compared by name:"
)
doc.add_paragraph("Classes — exact name match.", style="List Bullet")
doc.add_paragraph(
    "Methods — within matched classes, checked with rename awareness: "
    "__init__ may map to setup or __call__, forward maps to __call__. "
    "Methods like reset_parameters are treated as always-inlined (Flax "
    "handles them via initialiser arguments). Private/helper methods "
    "within a class that has __call__ are treated as legitimately inlined.",
    style="List Bullet",
)
doc.add_paragraph(
    "Functions — a PyTorch function is also considered matched if it was "
    "promoted to a class in the output.",
    style="List Bullet",
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Formula:  ")
run.bold = True
p.add_run("score = (matched_classes + matched_methods + matched_functions) "
          "/ (total_classes + total_methods + total_functions) * 100")

doc.add_heading("9.2 Correctness Score (LLM-Based)", level=2)
doc.add_paragraph(
    "The ValidationAgent is run against the source and output. Deviations "
    "are filtered for known false positives (low-severity method_placement, "
    "missing_component, and dropped_feature are excluded as they represent "
    "legitimate Flax idioms)."
)
doc.add_paragraph(
    "Each remaining deviation contributes a penalty based on severity:"
)

t5 = doc.add_table(rows=4, cols=2, style="Light Shading Accent 1")
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Severity", "Penalty"]):
    t5.rows[0].cells[i].text = h
    for r in t5.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for row_idx, (s, p_val) in enumerate([
    ("High", "5"), ("Medium", "3"), ("Low", "1"),
], 1):
    t5.rows[row_idx].cells[0].text = s
    t5.rows[row_idx].cells[1].text = p_val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Formula:  ")
run.bold = True
p.add_run("budget = total_components * 3  (medium severity weight)")
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run("             score = max(0,  (1 - penalty / budget) * 100)")
doc.add_paragraph()
doc.add_paragraph(
    "The budget scales with codebase size, so a large repository with "
    "150+ components is not unfairly penalised compared to a small one. "
    "A medium-severity deviation on every single component yields 0%. "
    "A high-severity deviation costs more than one component's budget "
    "(5 > 3), appropriately penalising severe issues."
)

doc.add_heading("9.3 Utility File Verification", level=2)
doc.add_paragraph(
    "If both merged_utils.py and the corresponding _utils_jax.py output "
    "exist, Step 5 runs the same completeness check on utility files: "
    "extract components via AST, compare by name, and compute a "
    "completeness score. The utility score is printed alongside the model "
    "score and saved to the JSON scorecard under the utils_completeness key."
)

doc.add_heading("9.4 Overall Score", level=2)
p = doc.add_paragraph()
run = p.add_run("Formula:  ")
run.bold = True
p.add_run("overall = (completeness + correctness) / 2")
doc.add_paragraph()
doc.add_paragraph(
    "Results are saved as verification_scorecard.json in the output "
    "directory, including full deviation details and utility completeness "
    "for post-mortem analysis."
)


# ══════════════════════════════════════════════════════════════════════
# 10. Agent Architecture
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("10. Agent Architecture", level=1)

doc.add_paragraph(
    "The conversion is orchestrated by four specialised agents, each "
    "with a single responsibility:"
)

t_agents = doc.add_table(rows=5, cols=3, style="Light Shading Accent 1")
t_agents.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Agent", "File", "Responsibility"]):
    t_agents.rows[0].cells[i].text = h
    for r in t_agents.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
agents = [
    ("PrimaryAgent", "primary_agent.py",
     "Top-level orchestrator: routes files, fills gaps, runs "
     "validate/repair loop"),
    ("ModelConversionAgent", "model_conversion_agent.py",
     "Converts nn.Module files using MODEL_CONVERSION_PROMPT with 16 "
     "Flax-specific rules"),
    ("SingleFileAgent", "single_file_agent.py",
     "Converts utility/non-model files using MIGRATE_MODULE_TO_JAX_PROMPT "
     "with general JAX patterns"),
    ("ValidationAgent", "validation_agent.py",
     "Detects faithfulness deviations (6 categories) and repairs them "
     "with RAG-augmented prompts"),
]
for row_idx, (agent, file, resp) in enumerate(agents, 1):
    t_agents.rows[row_idx].cells[0].text = agent
    t_agents.rows[row_idx].cells[1].text = file
    t_agents.rows[row_idx].cells[2].text = resp

doc.add_paragraph()
doc.add_paragraph(
    "All agents share a RAGAgent instance for retrieving reference patterns. "
    "The RAGAgent wraps an EmbeddingAgent (Gemini embedding-001) and the "
    "SQLite vector database."
)


# ══════════════════════════════════════════════════════════════════════
# 11. Architecture Diagram
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("11. Architecture Diagram", level=1)

diagram = doc.add_paragraph()
diagram.paragraph_format.space_before = Pt(6)
diagram.paragraph_format.space_after = Pt(6)
run = diagram.add_run(
    "PyTorch Repository\n"
    "       |\n"
    "       v\n"
    " [Step 1: Clone]\n"
    "       |\n"
    "       v\n"
    " [Step 2: Index] ---------> RAG Vector DB (48 docs, embedding-001)\n"
    "       |                              |\n"
    "       v                              |\n"
    " [Step 3a: Merge Models]              |  (hybrid per-component retrieval)\n"
    "       |                              |\n"
    "       |--- model files               |\n"
    "       |    (nn.Module)               |\n"
    "       v                              |\n"
    " [Step 3b: Discover & Merge Utils]    |\n"
    "       |                              |\n"
    "       |--- BFS from model imports    |\n"
    "       |--- classify (5 categories)   |\n"
    "       |--- filter & topo-sort        |\n"
    "       |                              |\n"
    "       v                              v\n"
    " merged_model.py  ---------> [Step 4: Convert Models]\n"
    " merged_utils.py  --|              |\n"
    "                    |         ModelConversionAgent\n"
    "                    |              |\n"
    "                    |         Fill Missing Components\n"
    "                    |         (Phase 1 + Phase 2)\n"
    "                    |              |\n"
    "                    |         Validate & Repair\n"
    "                    |         (up to 3 iters)\n"
    "                    |              |\n"
    "                    |              v\n"
    "                    |         <repo>_jax.py\n"
    "                    |\n"
    "                    +------> [Step 4: Convert Utils]\n"
    "                                   |\n"
    "                              SingleFileAgent\n"
    "                                   |\n"
    "                              Fill Missing Components\n"
    "                                   |\n"
    "                                   v\n"
    "                              <repo>_utils_jax.py\n"
    "                                   |\n"
    "       ,----------------------------'\n"
    "       v\n"
    " [Step 5: Verify]\n"
    "       |\n"
    "       |--- Model: Completeness + Correctness\n"
    "       |--- Utils: Completeness\n"
    "       |\n"
    "       v\n"
    " verification_scorecard.json"
)
run.font.name = "Consolas"
run.font.size = Pt(9)


# ══════════════════════════════════════════════════════════════════════
# 12. Data Flow Summary
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("12. Data Flow Summary", level=1)

t_flow = doc.add_table(rows=8, cols=3, style="Light Shading Accent 1")
t_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Stage", "Input", "Output"]):
    t_flow.rows[0].cells[i].text = h
    for r in t_flow.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
flows = [
    ("Step 1: Clone", "Repository URL", "Local clone directory"),
    ("Step 2: Index", "rag/sources/*.py", "~/rag_store.db"),
    ("Step 3a: Merge Models", "Cloned repo .py files", "merged_model.py"),
    ("Step 3b: Merge Utils", "Model file import graph", "merged_utils.py"),
    ("Step 4: Convert Models", "merged_model.py + RAG DB", "<repo>_jax.py"),
    ("Step 4: Convert Utils", "merged_utils.py + RAG DB", "<repo>_utils_jax.py"),
    ("Step 5: Verify", "Source + output files", "verification_scorecard.json"),
]
for row_idx, (stage, inp, out) in enumerate(flows, 1):
    t_flow.rows[row_idx].cells[0].text = stage
    t_flow.rows[row_idx].cells[1].text = inp
    t_flow.rows[row_idx].cells[2].text = out


# ══════════════════════════════════════════════════════════════════════
# 13. Design Decisions
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("13. Key Design Decisions", level=1)

decisions = [
    ("Separate model and utility merges",
     "Utility files are merged into merged_utils.py, not mixed into "
     "merged_model.py. This keeps the model conversion path unchanged "
     "and makes utility handling purely additive."),
    ("SingleFileAgent for utilities",
     "Utility files are converted with SingleFileAgent, not "
     "ModelConversionAgent, because they contain no nn.Module subclasses. "
     "The model-specific conversion rules (compact decorator, setup vs "
     "__call__) do not apply."),
    ("Re-export __init__.py files skipped",
     "init_reexport files contain only import statements that are already "
     "inlined by the merge process. Including them would add duplicate "
     "code."),
    ("CUDA kernel loaders skipped",
     "Files that use load()/load_inline() to compile .cu/.cpp custom ops "
     "have no JAX equivalent. However, autograd.Function files that wrap "
     "these kernels are kept because they often have a Python fallback "
     "implementation worth converting."),
    ("Utility discovery seeded from final model file list",
     "The BFS starts from the required model files (after filtering and "
     "dependency tracing), not from all model files. This ensures only "
     "utilities actually needed by the included models are discovered."),
    ("Iterative repair with early exit",
     "The validate-repair loop runs at most 3 iterations and exits early "
     "if the deviation count does not decrease. This prevents infinite "
     "loops when the LLM introduces new issues while fixing old ones."),
    ("Ratio-based correctness scoring",
     "The correctness budget scales with codebase size "
     "(components x medium_weight), ensuring large repositories are not "
     "unfairly penalised compared to small ones."),
]
for title_text, desc in decisions:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc)


# ══════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, "MaxCode_Pipeline_Reference.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
